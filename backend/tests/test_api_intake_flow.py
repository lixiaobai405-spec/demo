from __future__ import annotations

from io import BytesIO
from pathlib import Path
import sys

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from app.db import session as db_session
from app.main import create_app
from app.models.assessment import Assessment  # noqa: F401
from app.models.intake_session import AssessmentIntakeSession  # noqa: F401
from app.services.intake_service import IntakeService


TEST_DB_PATH = Path(__file__).resolve().parent / "test_intake_flow.db"


@pytest.fixture()
def client(monkeypatch: pytest.MonkeyPatch) -> TestClient:
    if TEST_DB_PATH.exists():
        TEST_DB_PATH.unlink()

    engine = create_engine(
        f"sqlite:///{TEST_DB_PATH.as_posix()}",
        connect_args={"check_same_thread": False},
    )
    testing_session_local = sessionmaker(
        bind=engine,
        autocommit=False,
        autoflush=False,
    )

    monkeypatch.setattr(db_session, "engine", engine)
    monkeypatch.setattr(db_session, "SessionLocal", testing_session_local)

    app = create_app()
    db_session.Base.metadata.create_all(bind=engine)
    db_session._migrate_generated_reports_table()
    with TestClient(app) as test_client:
        register_response = test_client.post(
            "/api/auth/register",
            json={
                "email": "student@example.com",
                "password": "password123",
                "display_name": "Test Student",
            },
        )
        assert register_response.status_code == 201
        token = register_response.json()["access_token"]
        test_client.headers.update({"Authorization": f"Bearer {token}"})
        yield test_client

    engine.dispose()
    if TEST_DB_PATH.exists():
        TEST_DB_PATH.unlink()


@pytest.fixture()
def confirmed_assessment_input() -> dict[str, str]:
    return {
        "company_name": "测试连锁零售企业",
        "industry": "零售",
        "company_size": "100-499人",
        "region": "华东",
        "annual_revenue_range": "5000万-1亿元",
        "core_products": "社区零售门店、会员运营与到家服务",
        "target_customers": "社区家庭用户、周边白领与会员客户",
        "current_challenges": "门店运营效率波动，会员复购不稳定",
        "ai_goals": "提升门店运营效率，增强会员复购",
        "available_data": "POS、会员系统、商品主数据",
        "notes": "计划先从单区域试点推进",
    }


def test_intake_import_returns_prefill_and_field_meta(client: TestClient) -> None:
    response = client.post(
        "/api/intake/import",
        json={
            "source_type": "markdown",
            "raw_content": """
# 企业课前输入

企业名称：测试连锁零售企业
所属行业：零售
企业规模：100-499人
所在区域：华东
年营收范围：5000万-1亿元
核心产品/服务：社区零售门店、会员运营与到家服务
目标客户：社区家庭用户、周边白领与会员客户

目前最大的挑战是门店运营效率波动，会员复购不稳定，店长培养周期长。
需要通过 AI 改进的关键点：
1）门店排班与补货计划依赖人工经验，高峰时段服务响应慢；
2）会员生命周期缺乏自动化运营，高价值客户流失无预警；
3）门店知识分散在老员工脑中，新人上手周期 3-6 个月。
希望通过 AI 稳定门店管理并提升复购，沉淀可复制的门店运营知识体系。
当前可用数据包括 POS、会员系统、商品主数据，另有巡店和客服记录。
            """,
        },
    )

    assert response.status_code == 200
    body = response.json()

    assert body["status"] == "parsed"
    assert body["source_type"] == "markdown"
    assert body["import_session_id"]
    assert body["assessment_prefill"]["company_name"] == "测试连锁零售企业"
    assert body["assessment_prefill"]["annual_revenue_range"] == "5000万-1亿元"
    assert body["field_meta"]["company_name"] == {
        "source_type": "raw",
        "status": "confirmed",
    }
    assert body["field_meta"]["ai_goals"]["source_type"] in ("raw", "inferred")


def test_intake_import_maps_markdown_heading_blocks_to_partial_prefill(client: TestClient) -> None:
    response = client.post(
        "/api/intake/import",
        json={
            "source_type": "markdown",
            "raw_content": """
## 企业名称
测试工业软件企业

## 所属行业
工业软件

## 年营收范围
3000万-5000万

## 核心产品/服务
设备数据采集平台
智能报表系统

## 当前经营/管理挑战
项目交付依赖人工经验，需求响应速度不稳定。
产品更新迭代周期长，客户反馈到版本发布平均需要 4-6 周。

## AI 改进方向
希望引入 AI 自动诊断生产设备异常，并根据历史案例推荐处理方案，
缩短现场工程师的排障时间，提升首次修复率。
            """,
        },
    )

    assert response.status_code == 200
    body = response.json()

    assert body["assessment_prefill"]["company_name"] == "测试工业软件企业"
    assert body["assessment_prefill"]["industry"] == "工业软件"
    assert body["assessment_prefill"]["annual_revenue_range"] == "3000万-5000万"
    assert body["assessment_prefill"]["core_products"] == "设备数据采集平台\n智能报表系统"
    assert "项目交付依赖人工经验" in (body["assessment_prefill"]["current_challenges"] or "")
    assert body["assessment_prefill"]["target_customers"] is None
    assert any("目标客户未识别" in warning for warning in body["warnings"])


def test_intake_import_can_create_assessment_after_confirmation(
    client: TestClient,
    confirmed_assessment_input: dict[str, str],
) -> None:
    import_response = client.post(
        "/api/intake/import",
        json={
            "source_type": "form",
            "structured_fields": {
                "company_name": "测试连锁零售企业",
                "industry": "零售",
                "region": "华东",
            },
        },
    )
    assert import_response.status_code == 200
    import_session_id = import_response.json()["import_session_id"]

    create_response = client.post(
        f"/api/intake/import/{import_session_id}/assessment",
        json={"confirmed_assessment_input": confirmed_assessment_input},
    )

    assert create_response.status_code == 201
    body = create_response.json()
    assert body["status"] == "confirmed"
    assert body["import_session_id"] == import_session_id
    assert body["assessment"]["company_name"] == confirmed_assessment_input["company_name"]

    assessment_id = body["assessment"]["id"]
    detail_response = client.get(f"/api/assessments/{assessment_id}")
    assert detail_response.status_code == 200
    assert detail_response.json()["assessment"]["industry"] == confirmed_assessment_input["industry"]


def test_get_intake_import_session_detail_returns_persisted_payload(client: TestClient) -> None:
    import_response = client.post(
        "/api/intake/import",
        json={
            "source_type": "markdown",
            "raw_content": """
企业名称：测试连锁零售企业
所属行业：零售
企业规模：100-499人
所在区域：华东
年营收范围：5000万-1亿元
当前经营/管理挑战：门店运营效率波动，会员复购不稳定，店长培养周期长，一线人员流动率高
希望通过 AI 达成的目标：提升门店运营效率，增强会员复购
当前可用数据/系统基础：POS、会员系统、商品主数据
另有备注：优先选择华东区域试点。
            """,
            "structured_fields": {
                "region": "华东",
            },
        },
    )
    assert import_response.status_code == 200
    import_session_id = import_response.json()["import_session_id"]

    detail_response = client.get(f"/api/intake/import/{import_session_id}")

    assert detail_response.status_code == 200
    body = detail_response.json()
    assert body["import_session_id"] == import_session_id
    assert body["raw_content"]
    assert body["structured_fields"] == {"region": "华东"}
    assert body["assessment_prefill"]["company_name"] == "测试连锁零售企业"
    assert body["field_meta"]["company_name"]["status"] == "confirmed"
    assert body["field_candidates"]["available_data"]["value"] == "POS、会员系统、商品主数据"
    assert body["created_assessment_id"] is None
    assert body["created_at"]
    assert body["updated_at"]


def test_intake_file_upload_supports_docx(client: TestClient) -> None:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("企业名称：上传测试企业")
    document.add_paragraph("所属行业：制造")
    document.add_paragraph("希望通过 AI 达成的目标：提升设备维护效率")
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/api/intake/import/file",
        files={
            "file": (
                "intake.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["source_type"] == "file"
    assert body["source_file"] == {
        "name": "intake.docx",
        "kind": "docx",
        "size_bytes": len(buffer.getvalue()),
    }
    assert body["assessment_prefill"]["company_name"] == "上传测试企业"
    assert body["assessment_prefill"]["industry"] == "制造"


def test_intake_file_upload_rejects_unsupported_type(client: TestClient) -> None:
    response = client.post(
        "/api/intake/import/file",
        files={
            "file": (
                "intake.xlsx",
                b"fake-binary",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )

    assert response.status_code == 415
    assert response.json()["detail"] == "仅支持上传 txt、md、markdown、pdf、docx 文件。"


def test_intake_service_extracts_pdf_text_via_pdf_reader(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def __init__(self, text: str):
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class FakePdfReader:
        def __init__(self, _: BytesIO):
            self.is_encrypted = False
            self.pages = [
                FakePage("企业名称：PDF 测试企业"),
                FakePage("所属行业：零售"),
            ]

    monkeypatch.setattr("app.services.intake_service.PdfReader", FakePdfReader)

    extracted, warnings = IntakeService()._extract_text_from_pdf(b"%PDF-1.4 fake")
    assert extracted == "企业名称：PDF 测试企业\n所属行业：零售"
    assert warnings == []


def test_intake_service_falls_back_to_pdf_ocr(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def extract_text(self) -> str:
            return ""

    class FakePdfReader:
        def __init__(self, _: BytesIO):
            self.is_encrypted = False
            self.pages = [FakePage()]

    monkeypatch.setattr("app.services.intake_service.PdfReader", FakePdfReader)
    monkeypatch.setattr(
        IntakeService,
        "_extract_text_from_pdf_via_ocr",
        lambda self, file_bytes, page_count: (
            "企业名称：OCR 企业\n所属行业：制造",
            ["检测到扫描版或低文本密度 PDF，已自动启用 OCR 识别。"],
        ),
    )

    extracted, warnings = IntakeService()._extract_text_from_pdf(b"%PDF-1.4 fake")
    assert extracted == "企业名称：OCR 企业\n所属行业：制造"
    assert any("自动启用 OCR" in warning for warning in warnings)


def test_intake_service_returns_specific_error_for_large_file(client: TestClient) -> None:
    response = client.post(
        "/api/intake/import/file",
        files={
            "file": (
                "intake.pdf",
                b"x" * (10 * 1024 * 1024 + 1),
                "application/pdf",
            )
        },
    )

    assert response.status_code == 413
    assert "上传文件过大" in response.json()["detail"]
    assert "10.0 MB" in response.json()["detail"]


def test_intake_import_rejects_duplicate_assessment_creation(
    client: TestClient,
    confirmed_assessment_input: dict[str, str],
) -> None:
    import_response = client.post(
        "/api/intake/import",
        json={
            "source_type": "text",
            "raw_content": "企业名称：测试连锁零售企业\n所属行业：零售",
        },
    )
    assert import_response.status_code == 200
    import_session_id = import_response.json()["import_session_id"]

    first_response = client.post(
        f"/api/intake/import/{import_session_id}/assessment",
        json={"confirmed_assessment_input": confirmed_assessment_input},
    )
    second_response = client.post(
        f"/api/intake/import/{import_session_id}/assessment",
        json={"confirmed_assessment_input": confirmed_assessment_input},
    )

    assert first_response.status_code == 201
    assert second_response.status_code == 409
    assert (
        second_response.json()["detail"]
        == "This intake import session has already created an assessment."
    )


def test_get_intake_import_session_detail_shows_created_assessment_id(
    client: TestClient,
    confirmed_assessment_input: dict[str, str],
) -> None:
    import_response = client.post(
        "/api/intake/import",
        json={
            "source_type": "text",
            "raw_content": "企业名称：测试连锁零售企业\n所属行业：零售",
        },
    )
    assert import_response.status_code == 200
    import_session_id = import_response.json()["import_session_id"]

    create_response = client.post(
        f"/api/intake/import/{import_session_id}/assessment",
        json={"confirmed_assessment_input": confirmed_assessment_input},
    )
    assert create_response.status_code == 201
    assessment_id = create_response.json()["assessment"]["id"]

    detail_response = client.get(f"/api/intake/import/{import_session_id}")

    assert detail_response.status_code == 200
    body = detail_response.json()
    assert body["status"] == "confirmed"
    assert body["created_assessment_id"] == assessment_id


def test_get_intake_import_session_detail_returns_404_for_unknown_id(client: TestClient) -> None:
    response = client.get("/api/intake/import/not-exists")

    assert response.status_code == 404
    assert response.json()["detail"] == "Intake import session not found."


def test_assessment_detail_claims_legacy_ownerless_intake_assessment(
    client: TestClient,
    confirmed_assessment_input: dict[str, str],
) -> None:
    import_response = client.post(
        "/api/intake/import",
        json={
            "source_type": "form",
            "structured_fields": {
                "company_name": "legacy-company",
                "industry": "retail",
            },
        },
    )
    assert import_response.status_code == 200
    import_session_id = import_response.json()["import_session_id"]

    create_response = client.post(
        f"/api/intake/import/{import_session_id}/assessment",
        json={"confirmed_assessment_input": confirmed_assessment_input},
    )
    assert create_response.status_code == 201
    assessment_id = create_response.json()["assessment"]["id"]

    with db_session.SessionLocal() as db:
        assessment = db.get(Assessment, assessment_id)
        session = db.get(AssessmentIntakeSession, import_session_id)
        assert assessment is not None
        assert session is not None
        assessment.user_id = None
        session.user_id = None
        db.add(assessment)
        db.add(session)
        db.commit()

    detail_response = client.get(f"/api/assessments/{assessment_id}")

    assert detail_response.status_code == 200
    assert detail_response.json()["assessment"]["id"] == assessment_id

    with db_session.SessionLocal() as db:
        assessment = db.get(Assessment, assessment_id)
        session = db.get(AssessmentIntakeSession, import_session_id)
        assert assessment is not None
        assert session is not None
        assert assessment.user_id is not None
        assert session.user_id == assessment.user_id
