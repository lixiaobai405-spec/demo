import json
import re
from io import BytesIO

from docx import Document
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.assessment import Assessment
from app.models.intake_session import AssessmentIntakeSession
from app.services.intake_llm_extractor import IntakeLLMExtractor
from app.schemas.assessment import AssessmentResponse
from app.schemas.intake import (
    ASSESSMENT_FIELD_NAMES,
    AssessmentPrefillDraft,
    IntakeCreateAssessmentRequest,
    IntakeCreateAssessmentResponse,
    IntakeFieldCandidate,
    IntakeFieldMeta,
    IntakeImportRequest,
    IntakeImportResponse,
    IntakeSourceFile,
    IntakeSessionDetailResponse,
)

FIELD_LABELS: dict[str, list[str]] = {
    "company_name": ["企业名称", "公司名称", "企业名", "company_name"],
    "industry": ["所属行业", "行业", "industry"],
    "company_size": ["企业规模", "公司规模", "员工人数", "company_size"],
    "region": ["所在区域", "所在地区", "区域", "region"],
    "annual_revenue_range": ["年营收范围", "营收范围", "年度营收", "annual_revenue_range"],
    "core_products": ["核心产品/服务", "核心产品", "产品服务", "主营业务", "core_products"],
    "target_customers": ["目标客户", "客户群体", "目标用户", "target_customers"],
    "current_challenges": ["当前经营/管理挑战", "当前挑战", "经营挑战", "管理挑战", "current_challenges"],
    "ai_goals": ["希望通过 AI 达成的目标", "AI 目标", "智能化目标", "ai_goals"],
    "available_data": ["当前可用数据/系统基础", "可用数据", "系统基础", "available_data"],
    "notes": ["其他补充说明", "补充说明", "备注", "notes"],
}

FIELD_KEYWORDS: dict[str, list[str]] = {
    "company_name": ["名称", "公司", "企业名", "单位"],
    "industry": ["行业", "领域", "产业"],
    "company_size": ["规模", "人数", "大小", "员工"],
    "region": ["区域", "地区", "城市", "省份", "省", "市", "地址"],
    "annual_revenue_range": ["营收", "收入", "年收", "营业额", "产值"],
    "core_products": ["产品", "服务", "业务"],
    "target_customers": ["客户", "用户", "消费者", "客群"],
    "current_challenges": ["挑战", "问题", "痛点", "困难", "瓶颈", "不足", "障碍"],
    "ai_goals": ["AI", "目标", "希望", "期望", "愿景", "想要", "打算", "降本", "增效", "提效", "自动化"],
    "available_data": ["数据", "系统", "ERP", "CRM", "POS", "台账", "平台", "信息化", "数字化"],
    "notes": ["备注", "补充", "其他", "说明", "附注", "额外"],
}

FIELD_DISPLAY_NAMES: dict[str, str] = {
    "company_name": "企业名称",
    "industry": "所属行业",
    "company_size": "企业规模",
    "region": "所在区域",
    "annual_revenue_range": "年营收范围",
    "core_products": "核心产品/服务",
    "target_customers": "目标客户",
    "current_challenges": "当前经营/管理挑战",
    "ai_goals": "希望通过 AI 达成的目标",
    "available_data": "当前可用数据/系统基础",
    "notes": "其他补充说明",
}

INFERENCE_RULES: dict[str, list[str]] = {
    "current_challenges": ["挑战", "瓶颈", "问题", "痛点", "困难", "不足", "障碍", "低效", "太慢", "滞后"],
    "ai_goals": ["ai", "目标", "希望", "提升", "降本", "增效", "复购", "增长", "自动化", "智能"],
    "available_data": ["数据", "系统", "pos", "erp", "crm", "会员", "订单", "台账", "excel"],
}

SUPPORTED_UPLOAD_EXTENSIONS = {
    ".txt": "txt",
    ".md": "markdown",
    ".markdown": "markdown",
    ".pdf": "pdf",
    ".docx": "docx",
}


class IntakeService:
    _rapidocr_engine = None
    _llm_extractor: IntakeLLMExtractor | None = None

    @property
    def llm_extractor(self) -> IntakeLLMExtractor:
        if self._llm_extractor is None:
            self._llm_extractor = IntakeLLMExtractor()
        return self._llm_extractor

    def get_session_detail(
        self,
        db: Session,
        import_session_id: str,
    ) -> IntakeSessionDetailResponse:
        session = db.get(AssessmentIntakeSession, import_session_id)
        if session is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Intake import session not found.",
            )

        parsed_payload = self._load_json_payload(session.parsed_payload)
        normalized_payload = self._load_json_payload(session.normalized_payload)

        return IntakeSessionDetailResponse(
            import_session_id=session.id,
            status=session.status,
            source_type=session.source_type,
            source_file=(
                IntakeSourceFile.model_validate(normalized_payload["source_file"])
                if normalized_payload.get("source_file")
                else None
            ),
            raw_content=session.raw_content,
            structured_fields=self._load_json_payload(session.structured_fields_payload),
            assessment_prefill=AssessmentPrefillDraft.model_validate(
                normalized_payload.get("assessment_prefill", {})
            ),
            field_meta={
                field_name: IntakeFieldMeta.model_validate(value)
                for field_name, value in normalized_payload.get("field_meta", {}).items()
            },
            field_candidates={
                field_name: IntakeFieldCandidate.model_validate(value)
                for field_name, value in parsed_payload.get("field_candidates", {}).items()
            },
            unmapped_notes=parsed_payload.get("unmapped_notes", []),
            warnings=parsed_payload.get("warnings", []),
            created_assessment_id=session.created_assessment_id,
            created_at=session.created_at.isoformat(),
            updated_at=session.updated_at.isoformat(),
        )

    def import_content(
        self,
        db: Session,
        payload: IntakeImportRequest,
        source_file: IntakeSourceFile | None = None,
        extra_warnings: list[str] | None = None,
    ) -> IntakeImportResponse:
        field_candidates = self._build_field_candidates(payload)
        field_meta = self._build_field_meta(field_candidates)
        warnings = [*(extra_warnings or []), *self._build_warnings(field_meta)]
        unmapped_notes = self._collect_unmapped_notes(
            raw_content=payload.raw_content or "",
            field_candidates=field_candidates,
        )
        assessment_prefill = AssessmentPrefillDraft(
            **{
                field_name: candidate.value
                for field_name, candidate in field_candidates.items()
            }
        )

        parsed_payload = {
            "field_candidates": {
                field_name: candidate.model_dump()
                for field_name, candidate in field_candidates.items()
            },
            "unmapped_notes": unmapped_notes,
            "warnings": warnings,
        }
        normalized_payload = {
            "assessment_prefill": assessment_prefill.model_dump(),
            "field_meta": {
                field_name: meta.model_dump() for field_name, meta in field_meta.items()
            },
            "source_file": source_file.model_dump() if source_file else None,
        }

        session = AssessmentIntakeSession(
            source_type=payload.source_type,
            raw_content=payload.raw_content,
            structured_fields_payload=json.dumps(payload.structured_fields, ensure_ascii=False),
            parsed_payload=json.dumps(parsed_payload, ensure_ascii=False),
            normalized_payload=json.dumps(normalized_payload, ensure_ascii=False),
            status="parsed",
        )
        db.add(session)
        db.commit()
        db.refresh(session)

        return IntakeImportResponse(
            import_session_id=session.id,
            status="parsed",
            source_type=payload.source_type,
            source_file=source_file,
            assessment_prefill=assessment_prefill,
            field_meta=field_meta,
            field_candidates=field_candidates,
            unmapped_notes=unmapped_notes,
            warnings=warnings,
        )

    async def import_file(
        self,
        db: Session,
        upload_file: UploadFile,
    ) -> IntakeImportResponse:
        file_name = (upload_file.filename or "").strip()
        if not file_name:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="上传文件缺少文件名。",
            )

        file_kind = self._detect_upload_file_kind(file_name)
        file_bytes = await upload_file.read()
        file_size = len(file_bytes)
        if file_size == 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="上传文件为空，请重新选择。",
            )
        max_upload_size_bytes = self._get_max_upload_size_bytes()
        if file_size > max_upload_size_bytes:
            too_large_status = (
                status.HTTP_413_CONTENT_TOO_LARGE
                if hasattr(status, "HTTP_413_CONTENT_TOO_LARGE")
                else status.HTTP_413_REQUEST_ENTITY_TOO_LARGE
            )
            raise HTTPException(
                status_code=too_large_status,
                detail=(
                    f"上传文件过大，当前文件为 {self._format_size_label(file_size)}，"
                    f"请控制在 {self._format_size_label(max_upload_size_bytes)} 以内。"
                ),
            )

        raw_content, extraction_warnings = self._extract_text_from_upload(
            file_kind,
            file_bytes,
        )
        if not raw_content.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=(
                    "未从上传文件中提取到可解析文本，请确认文件内容清晰可读。"
                    "若为扫描版 PDF，请检查 OCR 依赖是否已安装。"
                ),
            )

        import_result = self.import_content(
            db,
            IntakeImportRequest(source_type="file", raw_content=raw_content),
            source_file=IntakeSourceFile(
                name=file_name,
                kind=file_kind,
                size_bytes=file_size,
            ),
            extra_warnings=extraction_warnings,
        )

        # Async: ingest to knowledge base (non-blocking)
        self._ingest_to_knowledge_base(raw_content, file_name)

        return import_result

    def _auto_create_assessment(
        self,
        db: Session,
        import_result: IntakeImportResponse,
    ) -> IntakeImportResponse:
        """Auto-create assessment from import prefill, set created_assessment_id on result."""
        try:
            created = self.create_assessment_from_session(
                db,
                import_result.import_session_id,
                IntakeCreateAssessmentRequest(
                    confirmed_assessment_input=import_result.assessment_prefill,
                ),
            )
            import_result.created_assessment_id = created.assessment.id
            import_result.status = "confirmed"
        except HTTPException:
            # If creation fails (e.g., already created), just return result as-is
            pass
        return import_result

    def create_assessment_from_session(
        self,
        db: Session,
        import_session_id: str,
        payload: IntakeCreateAssessmentRequest,
    ) -> IntakeCreateAssessmentResponse:
        session = db.get(AssessmentIntakeSession, import_session_id)
        if session is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Intake import session not found.",
            )
        if session.created_assessment_id:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="This intake import session has already created an assessment.",
            )

        assessment = Assessment(**payload.confirmed_assessment_input.model_dump())
        db.add(assessment)
        db.flush()

        session.status = "confirmed"
        session.created_assessment_id = assessment.id
        db.add(session)
        db.commit()
        db.refresh(assessment)

        return IntakeCreateAssessmentResponse(
            import_session_id=session.id,
            status="confirmed",
            assessment=AssessmentResponse.model_validate(assessment, from_attributes=True),
        )

    def _build_field_candidates(
        self,
        payload: IntakeImportRequest,
    ) -> dict[str, IntakeFieldCandidate]:
        candidates: dict[str, IntakeFieldCandidate] = {}
        raw_content = payload.raw_content or ""

        # Step 0: LLM extraction (covers all 11 fields at once)
        llm_extracted = self._try_llm_extraction(raw_content)

        # Step 1: Structured fields from form input (always trusted)
        for field_name, value in payload.structured_fields.items():
            if field_name in ASSESSMENT_FIELD_NAMES and value:
                candidates[field_name] = IntakeFieldCandidate(
                    value=value,
                    source="原文",
                    confidence="high",
                    evidence=f"结构化输入：{value}",
                )

        # Step 2: Fill gaps from LLM extraction
        for field_name in ASSESSMENT_FIELD_NAMES:
            if field_name in candidates:
                continue
            llm_value = llm_extracted.get(field_name, "")
            if llm_value:
                candidates[field_name] = IntakeFieldCandidate(
                    value=llm_value,
                    source="LLM 提取",
                    confidence="high",
                    evidence=llm_value,
                )

        # Step 3: Regex-based direct extraction for remaining gaps
        for field_name in ASSESSMENT_FIELD_NAMES:
            if field_name in candidates:
                continue
            direct_match = self._extract_direct_field_value(raw_content, field_name)
            if direct_match:
                candidates[field_name] = IntakeFieldCandidate(
                    value=direct_match,
                    source="原文",
                    confidence="high",
                    evidence=direct_match,
                )

        # Step 4: Keyword inference for special fields
        for field_name in ("current_challenges", "ai_goals", "available_data"):
            if field_name in candidates:
                continue
            inferred_value = self._infer_from_raw_content(raw_content, field_name)
            if inferred_value:
                candidates[field_name] = IntakeFieldCandidate(
                    value=inferred_value,
                    source="推断",
                    confidence="medium",
                    evidence=inferred_value,
                )

        return candidates

    def _try_llm_extraction(self, raw_text: str) -> dict[str, str]:
        """Try LLM extraction; return empty dict on failure."""
        if not raw_text or not raw_text.strip():
            return {}
        try:
            return self.llm_extractor.extract(raw_text)
        except Exception:
            return {}

    @staticmethod
    def _ingest_to_knowledge_base(raw_text: str, file_name: str) -> None:
        """Ingest uploaded document into RAG knowledge base (fire-and-forget)."""
        if not settings.rag_enabled:
            return
        try:
            from app.rag.embeddings import EmbeddingManager
            from app.rag.vector_store import VectorStore
            from app.rag.chunker import DocumentChunker
            from app.rag.retriever import RAGRetriever

            retriever = RAGRetriever(rag_enabled=True)
            result = retriever.ingest_document(
                text=raw_text,
                source_file=file_name,
                metadata={
                    "source_type": "user_upload",
                    "file_name": file_name,
                },
            )
            import logging
            logger = logging.getLogger(__name__)
            logger.info(
                "Knowledge base ingestion: %s — chunks=%s",
                result.get("status"),
                result.get("chunks_added", 0),
            )
        except Exception as exc:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning("Knowledge base ingestion failed: %s", exc)

    def _build_field_meta(
        self,
        field_candidates: dict[str, IntakeFieldCandidate],
    ) -> dict[str, IntakeFieldMeta]:
        field_meta: dict[str, IntakeFieldMeta] = {}
        for field_name in ASSESSMENT_FIELD_NAMES:
            candidate = field_candidates.get(field_name)
            if candidate is None:
                field_meta[field_name] = IntakeFieldMeta(
                    source_type="missing",
                    status="needs_user_input",
                )
                continue

            if candidate.source == "原文":
                field_meta[field_name] = IntakeFieldMeta(
                    source_type="raw",
                    status="confirmed",
                )
                continue

            field_meta[field_name] = IntakeFieldMeta(
                source_type="inferred",
                status="needs_user_confirmation",
            )
        return field_meta

    def _build_warnings(self, field_meta: dict[str, IntakeFieldMeta]) -> list[str]:
        warnings: list[str] = []
        for field_name, meta in field_meta.items():
            if meta.status == "needs_user_input":
                warnings.append(f"{FIELD_DISPLAY_NAMES[field_name]}未识别，请用户补充。")
            elif meta.status == "needs_user_confirmation":
                warnings.append(f"{FIELD_DISPLAY_NAMES[field_name]}为系统推断，请用户确认。")
        return warnings

    def _extract_direct_field_value(self, raw_content: str, field_name: str) -> str | None:
        """Match inline labels and markdown-style heading blocks."""
        if not raw_content.strip():
            return None

        labels = FIELD_LABELS.get(field_name, [])
        if not labels:
            return None

        lines = [self._normalize_intake_line(line) for line in raw_content.splitlines()]
        scored: list[tuple[int, int, str]] = []  # (score, value_length, value)

        for index, cleaned in enumerate(lines):
            if not cleaned:
                continue
            lowered = cleaned.lower()
            hits = sum(1 for label in labels if label.lower() in lowered)
            if hits == 0:
                continue

            inline_value = self._extract_inline_field_value(cleaned, labels)
            if inline_value:
                scored.append((hits + 3, len(inline_value), inline_value))
                continue

            if self._looks_like_field_header(cleaned, labels):
                block_value = self._collect_following_block_value(lines, index)
                if block_value:
                    scored.append((hits + 2, len(block_value), block_value))
                continue

            scored.append((hits, len(cleaned), cleaned))

        if not scored:
            return None

        scored.sort(key=lambda x: (x[0], x[1]), reverse=True)
        return scored[0][2]

    def _infer_from_raw_content(self, raw_content: str, field_name: str) -> str | None:
        if not raw_content.strip():
            return None

        lines = [self._normalize_intake_line(line) for line in raw_content.splitlines()]
        keywords = INFERENCE_RULES[field_name]
        for line in lines:
            normalized_line = line.strip()
            if not normalized_line:
                continue
            lowered = normalized_line.lower()
            if any(keyword in lowered for keyword in keywords):
                return normalized_line
        return None

    def _collect_unmapped_notes(
        self,
        raw_content: str,
        field_candidates: dict[str, IntakeFieldCandidate],
    ) -> list[str]:
        if not raw_content.strip():
            return []

        used_evidence: set[str] = set()
        for candidate in field_candidates.values():
            for line in candidate.evidence.splitlines():
                normalized_line = self._normalize_intake_line(line)
                if normalized_line:
                    used_evidence.add(normalized_line)
        unmapped_notes: list[str] = []
        for line in raw_content.splitlines():
            normalized_line = self._normalize_intake_line(line)
            if not normalized_line:
                continue
            if normalized_line in used_evidence:
                continue
            if self._looks_like_any_field_header(normalized_line):
                continue
            unmapped_notes.append(normalized_line)

        return unmapped_notes[:10]

    def _normalize_intake_line(self, line: str) -> str:
        return re.sub(r"^\s*(?:[-*+•]\s+|\d+[.)]\s+|#{1,6}\s+|>\s*)", "", line).strip()

    def _extract_inline_field_value(self, line: str, labels: list[str]) -> str | None:
        parts = re.split(r"[：:＝=]", line, maxsplit=1)
        if len(parts) != 2:
            return None

        prefix, suffix = parts
        normalized_prefix = prefix.strip().lower()
        if not any(label.lower() in normalized_prefix for label in labels):
            return None

        value = suffix.strip()
        return value or None

    def _looks_like_field_header(self, line: str, labels: list[str]) -> bool:
        normalized_line = re.sub(r"[：:＝=]\s*$", "", line).strip()
        if not normalized_line:
            return False

        lowered = normalized_line.lower()
        if not any(label.lower() in lowered for label in labels):
            return False

        if re.search(r"[，。；;,.!?？]", normalized_line):
            return False

        return len(normalized_line) <= 24

    def _looks_like_any_field_header(self, line: str) -> bool:
        for labels in FIELD_LABELS.values():
            if self._extract_inline_field_value(line, labels):
                return True
            if self._looks_like_field_header(line, labels):
                return True
        return False

    def _collect_following_block_value(self, lines: list[str], start_index: int) -> str | None:
        collected: list[str] = []

        for line in lines[start_index + 1:]:
            if not line:
                if collected:
                    break
                continue

            if self._looks_like_any_field_header(line):
                break

            collected.append(line)

        value = "\n".join(collected).strip()
        return value or None

    def _detect_upload_file_kind(self, file_name: str) -> str:
        normalized_name = file_name.lower()
        for extension, kind in SUPPORTED_UPLOAD_EXTENSIONS.items():
            if normalized_name.endswith(extension):
                return kind

        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="仅支持上传 txt、md、markdown、pdf、docx 文件。",
        )

    def _extract_text_from_upload(
        self,
        file_kind: str,
        file_bytes: bytes,
    ) -> tuple[str, list[str]]:
        if file_kind in {"txt", "markdown"}:
            return self._decode_text_bytes(file_bytes), []
        if file_kind == "pdf":
            return self._extract_text_from_pdf(file_bytes)
        if file_kind == "docx":
            return self._extract_text_from_docx(file_bytes), []
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="暂不支持该文件类型。",
        )

    def _decode_text_bytes(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="文本文件编码无法识别，请转为 UTF-8 后重试。",
        )

    def _extract_text_from_pdf(self, file_bytes: bytes) -> tuple[str, list[str]]:
        try:
            reader = PdfReader(BytesIO(file_bytes))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception as exc:  # pragma: no cover - third-party behavior
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail="PDF 已加密，暂不支持解析。",
                    ) from exc

            parts = [(page.extract_text() or "").strip() for page in reader.pages]
            direct_text = "\n".join(part for part in parts if part).strip()
            warnings: list[str] = []

            if not self._should_try_pdf_ocr(direct_text):
                return direct_text, warnings

            try:
                ocr_text, ocr_warnings = self._extract_text_from_pdf_via_ocr(
                    file_bytes=file_bytes,
                    page_count=len(reader.pages),
                )
            except RuntimeError as exc:
                if direct_text:
                    warnings.append(f"{exc}，当前先返回 PDF 直读文本，结果可能不完整。")
                    return direct_text, warnings
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=str(exc),
                ) from exc

            if ocr_text:
                merged_text = self._merge_extracted_text(direct_text, ocr_text)
                warnings.extend(ocr_warnings)
                return merged_text, warnings

            if direct_text:
                return direct_text, warnings

            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="PDF 未提取到可解析文本，OCR 也未识别出有效内容，请确认扫描件清晰度。",
            )
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="PDF 解析失败，请确认文件未损坏且包含可复制文本。",
            ) from exc

    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        try:
            document = Document(BytesIO(file_bytes))
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="DOCX 解析失败，请确认文件未损坏。",
            ) from exc

        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ).strip()
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _get_max_upload_size_bytes(self) -> int:
        return max(settings.intake_max_upload_size_mb, 1) * 1024 * 1024

    def _should_try_pdf_ocr(self, extracted_text: str) -> bool:
        if not settings.intake_pdf_ocr_enabled:
            return False
        meaningful_chars = len(re.sub(r"\s+", "", extracted_text or ""))
        return meaningful_chars == 0

    def _extract_text_from_pdf_via_ocr(
        self,
        file_bytes: bytes,
        page_count: int,
    ) -> tuple[str, list[str]]:
        engine = self._get_rapidocr_engine()
        warnings = ["检测到扫描版或低文本密度 PDF，已自动启用 OCR 识别。"]
        png_pages, truncated = self._render_pdf_pages_as_png_bytes(
            file_bytes=file_bytes,
            max_pages=max(settings.intake_pdf_ocr_max_pages, 1),
        )
        if truncated:
            warnings.append(
                f"该 PDF 共 {page_count} 页，本次 OCR 仅处理前 {len(png_pages)} 页以控制解析时长。"
            )

        ocr_parts: list[str] = []
        for page_bytes in png_pages:
            page_text = self._extract_text_from_ocr_image_bytes(engine, page_bytes)
            if page_text:
                ocr_parts.append(page_text)

        return "\n".join(ocr_parts).strip(), warnings

    def _render_pdf_pages_as_png_bytes(
        self,
        file_bytes: bytes,
        max_pages: int,
    ) -> tuple[list[bytes], bool]:
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError(
                "当前环境缺少 PDF OCR 渲染依赖 `PyMuPDF`，无法解析扫描版 PDF。"
            ) from exc

        try:
            document = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as exc:
            raise RuntimeError("扫描版 PDF 渲染失败，请确认文件未损坏。") from exc

        try:
            total_pages = len(document)
            render_count = min(total_pages, max_pages)
            page_images: list[bytes] = []
            matrix = fitz.Matrix(2, 2)

            for page_index in range(render_count):
                page = document.load_page(page_index)
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                page_images.append(pixmap.tobytes("png"))

            return page_images, total_pages > render_count
        finally:
            document.close()

    def _get_rapidocr_engine(self):
        if IntakeService._rapidocr_engine is not None:
            return IntakeService._rapidocr_engine

        try:
            from rapidocr_onnxruntime import RapidOCR
        except ImportError as exc:
            raise RuntimeError(
                "当前环境缺少 OCR 依赖 `rapidocr_onnxruntime`，无法解析扫描版 PDF。"
            ) from exc

        try:
            IntakeService._rapidocr_engine = RapidOCR()
        except Exception as exc:
            raise RuntimeError(
                "OCR 引擎初始化失败，请检查模型下载网络或本地缓存配置。"
            ) from exc

        return IntakeService._rapidocr_engine

    def _extract_text_from_ocr_image_bytes(self, engine, image_bytes: bytes) -> str:
        try:
            result = engine(image_bytes)
        except Exception as exc:
            raise RuntimeError("OCR 识别失败，请稍后重试或改用文本粘贴导入。") from exc

        if hasattr(result, "txts"):
            texts = getattr(result, "txts", ()) or ()
            return "\n".join(text.strip() for text in texts if isinstance(text, str) and text.strip())

        if isinstance(result, tuple) and result:
            primary = result[0]
            if isinstance(primary, (list, tuple)):
                collected: list[str] = []
                for item in primary:
                    if isinstance(item, (list, tuple)) and len(item) >= 2 and isinstance(item[1], str):
                        text = item[1].strip()
                        if text:
                            collected.append(text)
                return "\n".join(collected)

        return ""

    def _merge_extracted_text(self, direct_text: str, ocr_text: str) -> str:
        direct = direct_text.strip()
        ocr = ocr_text.strip()
        if not direct:
            return ocr
        if not ocr:
            return direct
        if ocr in direct:
            return direct
        if direct in ocr:
            return ocr
        return f"{direct}\n{ocr}"

    def _format_size_label(self, size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} B"
        if size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        return f"{size_bytes / (1024 * 1024):.1f} MB"

    def _load_json_payload(self, payload: str | None) -> dict:
        if not payload:
            return {}
        loaded = json.loads(payload)
        return loaded if isinstance(loaded, dict) else {}
