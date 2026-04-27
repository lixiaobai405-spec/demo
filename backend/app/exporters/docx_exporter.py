from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.schemas.assessment import (
    ReportCardData,
    ReportData,
    ReportSectionData,
    ReportTableData,
)


class DocxExporter:
    def export(
        self,
        report_data: ReportData,
        target_path: Path,
        metadata: dict | None = None,
    ) -> Path:
        document = Document()
        document.core_properties.title = report_data.title

        self._configure_styles(document)
        self._render_header(document, report_data)
        self._render_metadata(document, metadata or {})

        for index, section in enumerate(report_data.sections, start=1):
            self._render_section(document, index, section)

        target_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(target_path)
        return target_path

    def _configure_styles(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Microsoft YaHei"
        normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.space_after = Pt(8)
        normal_style.paragraph_format.line_spacing = 1.35

    def _render_header(self, document: Document, report_data: ReportData) -> None:
        title = document.add_heading(report_data.title, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = document.add_paragraph(report_data.subtitle)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        meta = document.add_paragraph()
        meta.add_run(f"企业：{report_data.company_name}    ").bold = True
        meta.add_run(f"行业：{report_data.industry}    ")
        meta.add_run(f"规模：{report_data.company_size}    ")
        meta.add_run(f"区域：{report_data.region}")

        revenue = document.add_paragraph()
        revenue.add_run("营收范围：").bold = True
        revenue.add_run(report_data.annual_revenue_range)

        score = document.add_paragraph()
        score.add_run("AI 就绪度评分：").bold = True
        score.add_run(str(report_data.ai_readiness_score))
        score.add_run(f"  |  {report_data.ai_readiness_summary}")

        document.add_paragraph(
            "本报告基于当前问卷、结构化诊断结果和模板规则生成，不使用自由写作式 LLM 深度生成。"
        )

    def _render_metadata(self, document: Document, metadata: dict) -> None:
        document.add_heading("生成元信息", level=1)
        document.add_paragraph(
            f"generation_mode: {metadata.get('generation_mode', 'template')}"
        )
        document.add_paragraph(
            f"used_llm: {'true' if metadata.get('used_llm') else 'false'}"
        )
        document.add_paragraph(
            f"used_rag: {'true' if metadata.get('used_rag') else 'false'}"
        )

        warnings = metadata.get("warnings") or []
        if warnings:
            document.add_paragraph("warnings:")
            for item in warnings:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph("warnings: none")

    def _render_section(
        self,
        document: Document,
        index: int,
        section: ReportSectionData,
    ) -> None:
        document.add_heading(f"{index}. {section.title}", level=1)
        document.add_paragraph(section.content)

        for bullet in section.bullets:
            document.add_paragraph(bullet, style="List Bullet")

        for card in section.cards:
            self._render_card(document, card)

        if section.table:
            self._render_table(document, section.table)

        if section.note:
            note = document.add_paragraph()
            note.add_run("备注：").bold = True
            note.add_run(section.note)

    def _render_card(self, document: Document, card: ReportCardData) -> None:
        document.add_heading(card.title, level=2)
        if card.subtitle:
            subtitle = document.add_paragraph()
            subtitle.add_run(card.subtitle).italic = True
        if card.highlight:
            highlight = document.add_paragraph()
            highlight.add_run("重点：").bold = True
            highlight.add_run(card.highlight)
        document.add_paragraph(card.content)
        for bullet in card.bullets:
            document.add_paragraph(bullet, style="List Bullet 2")

    def _render_table(self, document: Document, table_data: ReportTableData) -> None:
        if not table_data.columns:
            return

        table = document.add_table(rows=1, cols=len(table_data.columns))
        table.style = "Table Grid"
        table.autofit = True

        header_cells = table.rows[0].cells
        for index, column in enumerate(table_data.columns):
            header_cells[index].text = column

        for row_values in table_data.rows:
            row_cells = table.add_row().cells
            for index, value in enumerate(row_values):
                row_cells[index].text = value

        document.add_paragraph()
